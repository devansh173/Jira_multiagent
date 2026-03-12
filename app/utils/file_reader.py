import os
import io


def extract_text(file) -> str:
    """
    Extract plain text from an uploaded file.
    Supports: .txt, .md, .pdf, .docx, .xlsx, .csv

    Returns extracted text as a string.
    Raises ValueError for unsupported file types.
    """
    filename  = file.filename.lower()
    extension = os.path.splitext(filename)[1]

    if extension in (".txt", ".md"):
        return _read_text(file)

    elif extension == ".pdf":
        return _read_pdf(file)

    elif extension == ".docx":
        return _read_docx(file)

    elif extension in (".xlsx", ".xls"):
        return _read_excel(file)

    elif extension == ".csv":
        return _read_csv(file)

    else:
        raise ValueError(
            f"Unsupported file type: {extension}. "
            "Supported types: .txt, .md, .pdf, .docx, .xlsx, .csv"
        )


# ---------------------------------------------------------------------------
# READERS
# ---------------------------------------------------------------------------

def _read_text(file) -> str:
    """Read plain text or markdown files."""
    return file.read().decode("utf-8", errors="ignore").strip()


def _read_pdf(file) -> str:
    """Extract text from a digital PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF support. Run: pip install pdfplumber")

    text   = []
    data   = file.read()
    source = io.BytesIO(data)

    with pdfplumber.open(source) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text.strip())

    return "\n\n".join(text)


def _read_docx(file) -> str:
    """Extract text from a Word document using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for Word support. Run: pip install python-docx")

    data     = file.read()
    source   = io.BytesIO(data)
    document = Document(source)
    lines    = []

    for para in document.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())

    # Also extract text from tables
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def _read_excel(file) -> str:
    """Extract text from an Excel file using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl is required for Excel support. Run: pip install openpyxl")

    data   = file.read()
    source = io.BytesIO(data)
    wb     = openpyxl.load_workbook(source, data_only=True)
    lines  = []

    for sheet in wb.worksheets:
        lines.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell) for cell in row if cell is not None]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def _read_csv(file) -> str:
    """Read a CSV file as plain text."""
    return file.read().decode("utf-8", errors="ignore").strip()


# ---------------------------------------------------------------------------
# TRUNCATE
# ---------------------------------------------------------------------------

def truncate_text(text: str, max_chars: int = 8000) -> str:
    """
    Truncate extracted text to avoid exceeding the model's context window.
    Adds a notice when truncated so the agent knows content was cut.
    """
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[... Document truncated to fit context window ...]"